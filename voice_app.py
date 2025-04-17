import streamlit as st
from openai import OpenAI
import json
from docx import Document as DocxDocument  # To avoid conflict with Word export
from PyPDF2 import PdfReader
from datetime import datetime
import re

# Initialize OpenAI client
client = OpenAI(api_key=st.secrets["OPENAI_API_KEY"])

# ---------------------- COMPLETE VOICE PROFILE ----------------------
VOICE_PROFILE = {
    "communication_style": {
        "Direct": "No fluff. Just the point. With a CTA.",
        "Encouraging": "Supportive and constructive. Kind.",
        "Playful": "Cheeky and casual. No characters or fantasy. Think cristmas cracker with more charm.",
        "Witty": "Sharp with a hilarious dry English style twist. The Office but only the laugh out loud parts.",
        "Professional": "Stick to clarity and credibility. Prioritize useful information over personality. Avoid slang, metaphors, and theatrics. Corporate with a CTA.",
        "Warm": "Friendly and human. Like a soft hug from your favourite Aunt.",
        "Bold": "Confident, strong statements. If you walk into a room everyone notices."
    },
    "content_format": {
        "Step-by-Step": "Give me a clear numbered sequence.",
        "Quick Summary": "Just the key points, fast. No waffle. Short sentences.",
        "Detailed Breakdown": "Explain clearly with depth, reasons, numbered lists and bullet points.",
        "Action List": "What do I do first and next? Give bullet points.",
        "Analytical": "Back it up with logic and reasons. Quotes from famous people where appropriate.",
        "Conversational": "Make it feel like a chat."
    },
    "generation": {
        "Gen Alpha (b.2013–2025)": "Immersed in tech. Intuitive and playful. Include many emojis.",
        "Gen Z (b.1997–2012)": "Fast, visual, and meme-fluent. Include emojis.",
        "Millennials (b.1990–1996)": "Digital-native. Likes social and gamified tone. Include a few emojis.",
        "Wise Millennials (b.1981–1989)": "Bridges analog and digital. Values clarity and feedback.",
        "Gen X (b.1965–1980)": "Independent and direct. Prefers practical and honest tone.",
        "Boomers (b.1946–1964)": "Structured and respectful. Clear value and reliability.",
        "Silent Generation (1928–1945)": "Formal, respectful, and rooted in tradition. Responds to clarity, courtesy, and structured messaging.",
        "Mixed": "Blend tone and rhythm across generations. Focus on clarity."
    },
    "length": {
        "Short": "Keep content under 100 words",
        "Medium": "100-200 word range",
        "Long": "200-500 word detailed content"
    }
}

# ---------------------- PROFILE MANAGEMENT ----------------------
def get_sidebar_profile():
    from PIL import Image, UnidentifiedImageError

    try:
        logo = Image.open("assets/logo2.png")
        st.sidebar.image(logo, width=100, use_container_width=False)
    except (FileNotFoundError, UnidentifiedImageError):
        st.sidebar.warning("Logo not found. Proceeding without it.")

    uploaded_file = st.sidebar.file_uploader(
        "Upload a reference doc. AI is not secure, don't upload secrets or scandals.", 
        type=["txt", "pdf", "docx"]
    )
    reference_text = extract_text_from_file(uploaded_file) if uploaded_file else ""

    with st.sidebar:
        st.header("Set The Vibe")

        # Add tone flair selector
        tone_flair = st.select_slider(
            "Tone Flair",
            options=["Nip", "Slash", "Blaze"],
            value="Slash",
            help="Choose how bold you want the writing to sound: Subtly edgy (Nip), sharp but stylish (Slash) or bold and direct (Blaze)."
        )

        ultra_direct = st.toggle(
            "Ultra-Direct Mode",
            value=False,
            help="Override all personality settings to deliver sharp, efficient content with no Blaze tone."
        )

       
        return {
            "communication_style": st.selectbox("Tone of voice", list(VOICE_PROFILE["communication_style"].keys()),
                              index=3,
                              help="Think, 'If I was my audience, how would I like to be spoken to?'"
            ),
            "content_format": st.selectbox("How should it look?", list(VOICE_PROFILE["content_format"].keys()),
                              index=1,
                              help="What are we going with today, detailed, chatty, action?"
            ),
            "generation": st.selectbox("Audience Style", list(VOICE_PROFILE["generation"].keys()),
                              index=7,
                              help="This adjusts tone pacing, references, and formality."
            ),
            "length": st.radio("Content Length", 
                             ["Short", "Medium", "Long"],
                             index=1
            ),  # Default to Medium
            "tone_flair": tone_flair,  # Include tone_flair in the returned profile
            "ultra_direct": ultra_direct,
            "reference_text": reference_text 
        }      

def validate_profile(profile):
    """Ensure all profile fields are selected"""
    missing = []
    required_fields = list(VOICE_PROFILE.keys()) + ["length"]
    for field in required_fields:
        if not profile.get(field):
            missing.append(f"Please select {field.replace('_', ' ').title()}")
    return missing

if "last_response" not in st.session_state:
    st.session_state.last_response = ""

def extract_text_from_file(uploaded_file):
    if uploaded_file.name.endswith(".txt"):
        return uploaded_file.read().decode("utf-8")

    elif uploaded_file.name.endswith(".pdf"):
        reader = PdfReader(uploaded_file)
        return "\n".join([page.extract_text() or "" for page in reader.pages])

    elif uploaded_file.name.endswith(".docx"):
        doc = DocxDocument(uploaded_file)
        return "\n".join([para.text for para in doc.paragraphs])

    return ""

def generate_filename(prompt: str, ext="docx") -> str:
    words = re.findall(r'\b\w+\b', prompt.lower())
    top_two = sorted(words[:2]) if len(words) >= 2 else ["content"]
    base = "".join(top_two)
    date_str = datetime.now().strftime("%d%b%Y")  # e.g. 15Apr2025
    return f"{base}_{date_str}.{ext}"

# ---------------------- CORE ENGINE ----------------------
def build_hidden_instructions(profile):
    tone_overrides = []

    ref_block = ""
    if profile.get("reference_text"):
        ref_block = "\n## Reference Material Provided:\n" + profile["reference_text"][:2000]
    
    if profile.get("ultra_direct", False):
        content = [
            "Ultra-Direct Mode is ON.",
            "Write as if you’re a real person who wants to instruct — quickly.",
            "Drop the charm. Avoid metaphors, intros, or creative phrasing.",
            "Be concise, direct, and blunt — with just enough human edge to not sound robotic.",
            "No warm-ups. No analogies. No fluff.",
            "Use full punctuation and capitalization. Do not mimic informal lowercase writing unless explicitly asked.",
            "",
            f"Content Format: {VOICE_PROFILE['content_format'][profile['content_format']]}",
            f"Generation: {VOICE_PROFILE['generation'][profile['generation']]}",
            f"Length: {VOICE_PROFILE['length'][profile['length']]}",
            ref_block
        ]
        return "\n".join(content)

    if profile["tone_flair"] == "Blaze" and profile["communication_style"] in ["Professional", "Direct"]:
        content = [
            "Blaze Mode — Executive edition.",
            "Tone must be bold, direct, and human. Skip metaphors, branded sign-offs, or dramatic flourishes.",
            "Keep sentences short. Prioritize frictionless clarity with a confident edge.",
            "No wordplay. No pep talk. This isn’t advertising — it’s communication.",
            "",
            f"Content Format: {VOICE_PROFILE['content_format'][profile['content_format']]}",
            f"Generation: {VOICE_PROFILE['generation'][profile['generation']]}",
            f"Length: {VOICE_PROFILE['length'][profile['length']]}",
            ref_block
        ]
        return "\n".join(content)
    
    if profile["communication_style"] in ["Professional", "Direct"]:
        tone_overrides.append(
            "Deliver bold, decisive statements. Skip the jokes, metaphors, or quirky analogies. Be assertive without being performative. No emojis, no theatrics — just charisma and clarity."
        )

    core_tone = [
        "You are Custom Content AI — a content generator with bite, style, and zero tolerance for corporate fluff.",
        "Your default tone is bold, modern, and irreverent. Think: texting a clever friend who's mildly distracted, but will absolutely roast you if you waste their time.",
        "",
        "## Core Tone Rules:",
        "Write like you're texting a mildly distracted friend — clear, casual, and charming.",
        "Avoid big words and formal tone — this isn’t a TED Talk or a bank chatbot.",
        "Clarity comes first, but don’t sacrifice personality. Think charm over polish.",
        "Stay human, stay cheeky, and never sound like LinkedIn on a Monday."
        "Do not create themes, characters, metaphors, or narrative devices unless explicitly requested. No unicorns or pirates. Avoid turning simple tasks into storytelling. Keep it grounded in real-world language and tone.",
    ]

    tone_flair = {
        "Nip": [
            "",
            "## Current Mood: Nip",
            "- Keep it clean, cut, and clever.",
            "- Say less, mean more — let the space between lines do some of the talking.",
            "- Dry wit wins. No sparkle, no fluff.",
            "- Use precision like a scalpel, not a spotlight.",
            "- If the line lingers in their mind later, you nailed it."
        ],
        "Slash": [
            "",
            "## Current Mood: Slash",
            "- Stylish. Smart. Intentional. Think editorial, not emotional.",
            "- If it cuts, it better look good doing it.",
            "- Irony is allowed, but never silly. Glossy and sharp, not shiny and loud.",
            "- Leave quotable lines — not punchlines.",
            "- Deliver like you're unfazed, not unimpressed."
        ],
        "Blaze": [
            "",
            "## Current Mood: Blaze",
            "- Confidence is the baseline. The tone should command, not beg.",
            "- Be bold, but don’t perform. Drop truths, not punchlines.",
            "- No hand-holding, no padding. Every line should feel like a decision.",
            "- Sarcasm is essential. Jokes are dryer than the sahara.",
            "- Cut the theatrics. You’re not on stage, you’re in charge."
        ]        
    }
 
    user_preferences = [
        "",
        "## User Preferences (flavor, not framework):",
        f"Communication Style: {VOICE_PROFILE['communication_style'][profile['communication_style']]}",
        f"Content Format: {VOICE_PROFILE['content_format'][profile['content_format']]}",
        f"Generation: {VOICE_PROFILE['generation'][profile['generation']]}",
        f"Length: {VOICE_PROFILE['length'][profile['length']]} - Be concise if short, thorough with bullet points if long",
        ref_block
    ]

    if profile.get("reference_text"):
        user_preferences.append(
            "\n## Reference Material Provided:\n" + profile["reference_text"][:2000]  # Truncate if long
        )


    return "\n".join(core_tone + tone_flair[profile["tone_flair"]] + tone_overrides + [""] + user_preferences)

# ---------------------- STREAMLIT APP ----------------------
st.set_page_config(page_title="Custom Content AI", layout="centered")
st.title("Blaze AI")
st.markdown(
    "Dry. Sarcastic. Maybe even usable.",
    unsafe_allow_html=True
)


# Instructions panel logic
if "instructions_shown" not in st.session_state:
    st.session_state.instructions_shown = True  # Show instructions by default

if st.session_state.instructions_shown:
    with st.expander("🧭 Start Here: How It Works"):
        st.markdown("""
        1. **Type your message idea below.**  
        Examples:
        - "An email to my team about the end-of-project deadline on Friday."
        - "A birthday card message for my dog groomer's aunt, Betty, in Spanish"  
        
        💡 **Be specific.** Blaze can’t read your mind. *Yet.*
        
        2. **Set the vibe** *(Top left arrow if you’re on mobile)*  
        - **Tone Flair** = how much of a menace Blaze is allowed to be
        - Tweak **voice, format & audience** — if you care about that kind of thing.  
        👈 **Sidebar’s** where the magic happens.
               
        3. **Optional:** Upload a file for extra context. Blaze will read the first 2000 words then get bored.
        
        4. **Hit** ➤  

       💡Blaze AI gives you one sharp response each time. **One message in, one response out.**
       
        5. **Download it** or lose it forever.

        6. **Want a remix?** Tweak the settings and hit ↩️ **Reuse Last Prompt.**

        #### TL;DR  
        👈 **Set your vibe in the sidebar  
        Type your idea  
        Hit ➤  
        Download it or lose it  
        Done!**
        
        """, unsafe_allow_html=True)

# Profile Management
profile = get_sidebar_profile()
st.session_state.profile = profile

# Chat interface and response rendering
if "last_prompt" not in st.session_state:
    st.session_state.last_prompt = ""
if "last_response" not in st.session_state:
    st.session_state.last_response = ""

if prompt := st.chat_input("Blaze it: What do you want written?"):
    st.session_state.instructions_shown = False
    st.session_state.last_prompt = prompt
    st.session_state.last_response = ""  # ✅ Clear last response before generating new one

    # Generate content
    system_msg = {"role": "system", "content": build_hidden_instructions(profile)}
    print("Calling OpenAI API with prompt:", prompt)
    
    with st.spinner("Stirring the sass…"):
        response = client.chat.completions.create(
            model="gpt-4",
            messages=[system_msg, {"role": "user", "content": prompt}]
        )

        # Save and display response
        content = response.choices[0].message.content
        st.session_state.last_response = content

        st.success("Here's your masterpiece.")

# --- Optional reuse of last prompt ---
if st.button("↩️ Reuse Last Prompt") and st.session_state.last_prompt:
    system_msg = {"role": "system", "content": build_hidden_instructions(profile)}
    print("🔁 Reusing last prompt:", st.session_state.last_prompt)
    response = client.chat.completions.create(
        model="gpt-4",
        messages=[system_msg, {"role": "user", "content": st.session_state.last_prompt}]
    )
    content = response.choices[0].message.content
    st.session_state.last_response = content

# --- Show last response if available ---
if st.session_state.last_response:
    st.markdown(f"**Original Request:** _{st.session_state.last_prompt}_")
    with st.chat_message("assistant"):
        st.markdown(st.session_state.last_response)

    filename_docx = generate_filename(st.session_state.last_prompt)
    filename_txt = filename_docx.replace(".docx", ".md")
    
    # Download as text
    st.download_button(
        label="📥 Download txt file",
        data=st.session_state.last_response,
        file_name=filename_txt
    )

    # Export as Word doc
    doc = DocxDocument()
    doc.add_heading("Your AI Writing", level=1)
    doc.add_paragraph(st.session_state.last_response)
    doc.save(filename_docx)

    with open(filename_docx, "rb") as file:
        st.download_button(
            label="📥 Download Word file",
            data=file,
            file_name=filename_docx,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

       

# Display conversation history
for msg in st.session_state.get("messages", []):
    with st.chat_message(msg["role"]):
        st.write(msg["content"])
